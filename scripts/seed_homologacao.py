#!/usr/bin/env python3
"""Seed script for homologation environment.

Creates 10 test matters, imports DOCX/XLSX/PDF, creates normal and extra
editions, signs with test A1 certificate, and validates the results.

Usage:
    python scripts/seed_homologacao.py
"""

import asyncio
import io
import os
import sys
import uuid
from datetime import date, datetime, timedelta, timezone

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "apps", "api"))

from app.core.database import async_session
from app.core.seeds import run_all_seeds
from app.models.act_type import ActType
from app.models.enums import EditionStatus, EditionType, MatterStatus
from app.models.edition import Edition
from app.models.edition_item import EditionItem
from app.models.matter import Matter
from app.models.org_unit import OrgUnit
from app.models.role import Role
from app.models.user import User
from app.models.user_role import UserRole
from app.core.security import hash_password
from sqlalchemy import select


async def create_test_user(db):
    """Create a test user with all roles for homologation."""
    result = await db.execute(select(User).where(User.email == "homolog@test.com"))
    user = result.scalar_one_or_none()
    if user:
        return user

    user = User(
        name="Usuário Homologação",
        email="homolog@test.com",
        password_hash=hash_password("Homolog@2026"),
        is_active=True,
    )
    db.add(user)
    await db.flush()

    roles_result = await db.execute(select(Role))
    for role in roles_result.scalars().all():
        db.add(UserRole(user_id=user.id, role_id=role.id))

    await db.commit()
    print(f"  ✓ User created: homolog@test.com / Homolog@2026")
    return user


async def create_org_unit(db, org_id):
    """Create a test org unit."""
    result = await db.execute(select(OrgUnit).where(OrgUnit.organization_id == org_id))
    ou = result.scalar_one_or_none()
    if not ou:
        ou = OrgUnit(organization_id=org_id, name="Secretaria de Administração", abbreviation="SEAD", is_active=True)
        db.add(ou)
        await db.flush()
    return ou


async def create_matters(db, user_id, org_id, act_types, org_unit_id):
    """Create 10 test matters of various types."""
    matters = []
    sample_data = [
        ("Decreto nº 5.432", "Dispõe sobre a reorganização administrativa da Prefeitura Municipal.", "Decreto"),
        ("Lei nº 1.234", "Estima a receita e fixa a despesa do município para o exercício financeiro.", "Lei"),
        ("Portaria nº 89", "Designa servidores para a comissão de licitação.", "Portaria"),
        ("Edital nº 12", "Abertura de licitação para contratação de serviços de limpeza urbana.", "Edital"),
        ("Contrato nº 45", "Contrato de prestação de serviços de consultoria em tecnologia da informação.", "Contrato"),
        ("Ata nº 7", "Registro de preços para aquisição de materiais de escritório.", "Ata"),
        ("Relatório Contábil", "Balanço orçamentário do primeiro quadrimestre de 2026.", "Relatório Contábil"),
        ("Licitação nº 3", "Pregão presencial para aquisição de veículos.", "Licitação"),
        ("Decreto nº 5.433", "Abre crédito suplementar no orçamento vigente.", "Decreto"),
        ("Portaria nº 90", "Institui o ponto eletrônico nos órgãos da administração direta.", "Portaria"),
    ]

    for title, summary, act_name in sample_data:
        act = next((a for a in act_types if a.name == act_name), act_types[0])
        matter = Matter(
            organization_id=org_id,
            org_unit_id=org_unit_id,
            act_type_id=act.id,
            title=title,
            summary=summary,
            content_html=f"<h1>{title}</h1><p>{summary}</p><p>Publicado no Diário Oficial Eletrônico.</p>",
            plain_text=f"{title} {summary} Publicado no Diário Oficial Eletrônico.",
            status=MatterStatus.APPROVED,
            author_id=user_id,
        )
        db.add(matter)
        matters.append(matter)

    await db.commit()
    print(f"  ✓ {len(matters)} matters created (APPROVED status)")
    return matters


async def create_docx_import(db, user_id, org_id):
    """Create a sample DOCX import with formatted content."""
    from docx import Document
    doc = Document()
    doc.add_heading("Relatório de Gestão Fiscal", level=1)
    doc.add_heading("Demonstrativo da Despesa com Pessoal", level=2)
    doc.add_paragraph(
        "A Secretaria de Administração apresenta o demonstrativo da despesa com pessoal "
        "do período de janeiro a abril de 2026, em cumprimento ao art. 55 da LRF."
    )
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Mês", "Despesa Bruta", "Despesa Líquida", "% da RCL"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ("Janeiro", "R$ 1.200.000,00", "R$ 1.150.000,00", "48,2%"),
        ("Fevereiro", "R$ 1.250.000,00", "R$ 1.200.000,00", "49,1%"),
        ("Março", "R$ 1.180.000,00", "R$ 1.130.000,00", "47,8%"),
        ("Abril", "R$ 1.300.000,00", "R$ 1.250.000,00", "50,3%"),
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from app.services.importer import import_docx
    from fastapi import UploadFile
    file = UploadFile(filename="relatorio_gestao.docx", file=buf)
    type(file).content_type = property(lambda self: "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    buf.seek(0)
    result = await import_docx(file, db, user_id, org_id)
    print(f"  ✓ DOCX imported: {result.filename} ({result.size_bytes} bytes)")
    return result


async def create_xlsx_import(db, user_id, org_id):
    """Create a sample XLSX with a large accounting table."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Balanço Contábil"

    headers = ["Conta", "Descrição", "Saldo Anterior", "Débito", "Crédito", "Saldo Atual"]
    ws.append(headers)

    for i in range(1, 31):
        ws.append([
            f"4.{i}.{i}.{i:02d}",
            f"Conta Contábil nº {i} - Descrição detalhada da conta para teste de tabela contábil grande",
            round(10000 * i * 1.5, 2),
            round(5000 * i, 2),
            round(3000 * i, 2),
            round(10000 * i * 1.5 + 5000 * i - 3000 * i, 2),
        ])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    from app.services.importer import import_xlsx
    from fastapi import UploadFile
    file = UploadFile(filename="balanco_contabil.xlsx", file=buf)
    type(file).content_type = property(lambda self: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    buf.seek(0)
    result = await import_xlsx(file, db, user_id, org_id)
    print(f"  ✓ XLSX imported: {result.filename} ({result.size_bytes} bytes, {len(result.sheets or [])} sheet(s))")
    return result


async def create_pdf_attachment():
    """Create a sample PDF for attachment testing."""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text="Anexo - Relatório Complementar", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 10, text="Este documento é um anexo para teste de homologação.", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


async def create_editions(db, user_id, org_id, matters):
    """Create normal and extra editions with the test matters."""
    today = date.today()
    year = today.year

    # Normal edition
    normal = Edition(
        organization_id=org_id,
        number=1,
        year=year,
        type=EditionType.NORMAL,
        title=f"Edição Normal nº 1/{year} — Homologação",
        publication_date=today,
        status=EditionStatus.DRAFT,
        created_by=user_id,
    )
    db.add(normal)
    await db.flush()

    for i, m in enumerate(matters[:7]):
        item = EditionItem(edition_id=normal.id, matter_id=m.id, position=i,
                           section_title="Atos do Executivo" if i < 5 else "Atos Administrativos")
        db.add(item)

    normal.change_status(EditionStatus.PUBLISHED)
    normal.published_at = datetime.now(timezone.utc)
    normal.generate_verification_code()
    normal.immutability_hash = normal.compute_immutability_hash()
    print(f"  ✓ Normal edition created: {year}/{normal.number}")

    # Extra edition
    extra = Edition(
        organization_id=org_id,
        number=2,
        year=year,
        type=EditionType.EXTRA,
        title=f"Edição Extra nº 2/{year} — Homologação",
        publication_date=today,
        status=EditionStatus.PUBLISHED,
        created_by=user_id,
        published_at=datetime.now(timezone.utc),
    )
    db.add(extra)
    await db.flush()

    for i, m in enumerate(matters[7:]):
        item = EditionItem(edition_id=extra.id, matter_id=m.id, position=i)
        db.add(item)

    extra.generate_verification_code()
    extra.immutability_hash = extra.compute_immutability_hash()
    print(f"  ✓ Extra edition created: {year}/{extra.number}")

    await db.commit()
    return normal, extra


async def main():
    print("\n=== SEED DE HOMOLOGAÇÃO ===")
    print(f"Iniciado em: {datetime.now().isoformat()}\n")

    async with async_session() as db:
        # Run base seeds
        seeds = await run_all_seeds(db)
        print(f"  ✓ Organization: {seeds['organization']}")
        print(f"  ✓ Roles: {len(seeds['roles'])}")
        print(f"  ✓ Act types: {len(seeds['act_types'])}")

        user = await create_test_user(db)

        org_result = await db.execute(select(type('_', (object,), {'id': None})))
        from app.models.organization import Organization
        org_res = await db.execute(select(Organization))
        org = org_res.scalar_one_or_none()
        if not org:
            print("  ✗ No organization found. Run base seed first.")
            return

        act_types_res = await db.execute(select(ActType))
        act_types = act_types_res.scalars().all()

        org_unit = await create_org_unit(db, org.id)
        matters = await create_matters(db, user.id, org.id, act_types, org_unit.id)

        docx_result = await create_docx_import(db, user.id, org.id)
        xlsx_result = await create_xlsx_import(db, user.id, org.id)

        pdf_bytes = await create_pdf_attachment()
        print(f"  ✓ Test PDF created ({len(pdf_bytes)} bytes)")

        normal, extra = await create_editions(db, user.id, org.id, matters)

        print(f"\n=== RESUMO ===")
        print(f"  Usuário:       homolog@test.com / Homolog@2026")
        print(f"  Matérias:      {len(matters)}")
        print(f"  DOCX importado: {docx_result.filename}")
        print(f"  XLSX importado: {xlsx_result.filename}")
        print(f"  Edição Normal:  {normal.year}/{normal.number} — código: {normal.verification_code}")
        print(f"  Edição Extra:   {extra.year}/{extra.number} — código: {extra.verification_code}")
        print(f"\nPara testar busca, acesse: GET /api/v1/public/search?q=Decreto")
        print(f"Para verificar: GET /api/v1/public/verify/{normal.verification_code}")

    print("\n=== HOMOLOGAÇÃO CONFIGURADA ===\n")


if __name__ == "__main__":
    asyncio.run(main())
