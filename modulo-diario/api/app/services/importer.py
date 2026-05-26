"""File import and conversion services."""

import csv
import io
import logging
import uuid
from datetime import datetime
from typing import Optional

from fastapi import UploadFile

from app.core.config import settings
from app.core.file_validator import validate_upload
from app.core.storage import storage
from app.models.file import File as FileModel
from app.providers.antivirus import get_virus_scanner
from app.services.pdf_utils import compute_hash

logger = logging.getLogger(__name__)


class ImportResult:
    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)


class SheetInfo:
    def __init__(self, name: str, columns: int, row_count: int):
        self.name = name
        self.columns = columns
        self.row_count = row_count


def _storage_path(ext: str) -> str:
    now = datetime.utcnow()
    return f"imports/{now.strftime('%Y/%m/%d')}/{uuid.uuid4()}{ext}"


async def _save_file(
    db_session,
    user_id: uuid.UUID,
    org_id: uuid.UUID,
    filename: str,
    content: bytes,
    ext: str,
    mime_type: str,
) -> FileModel:
    path = _storage_path(ext)
    await storage.store(path, content)
    file_hash = compute_hash(content)
    file_record = FileModel(
        organization_id=org_id,
        filename=filename,
        mime_type=mime_type,
        size_bytes=len(content),
        storage_path=path,
        storage_bucket=settings.MINIO_BUCKET,
        hash=file_hash,
        uploaded_by=user_id,
        is_temp=False,
    )
    db_session.add(file_record)
    await db_session.flush()
    return file_record


# ── DOCX ─────────────────────────────────────────────────────────────────────


async def import_docx(
    file: UploadFile,
    db_session,
    user_id: uuid.UUID,
    org_id: uuid.UUID,
) -> ImportResult:
    ext, content = await validate_upload(file)
    scanner = get_virus_scanner()
    scan = await scanner.scan(content, file.filename or "document.docx")
    if not scan.clean:
        raise ValueError(f"Virus detected: {scan.message}")

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    html_parts: list[str] = []

    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ""
        text = _paragraph_to_html(para)
        if not text.strip():
            continue
        if style.startswith("heading 1"):
            html_parts.append(f"<h1>{text}</h1>")
        elif style.startswith("heading 2"):
            html_parts.append(f"<h2>{text}</h2>")
        elif style.startswith("heading 3"):
            html_parts.append(f"<h3>{text}</h3>")
        else:
            html_parts.append(f"<p>{text}</p>")

    for table in doc.tables:
        html_parts.append(_table_to_html(table))

    content_html = "\n".join(html_parts)
    from app.core.html_sanitizer import sanitize_html
    content_html = sanitize_html(content_html)

    file_record = await _save_file(
        db_session, user_id, org_id,
        file.filename or "document.docx", content, ext,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    return ImportResult(
        file_id=file_record.id,
        filename=file_record.filename,
        size_bytes=file_record.size_bytes,
        hash=file_record.hash,
        mime_type=file_record.mime_type,
        content_html=content_html,
    )


def _paragraph_to_html(para) -> str:
    text = ""
    for run in para.runs:
        t = run.text
        if run.bold:
            t = f"<strong>{t}</strong>"
        if run.italic:
            t = f"<em>{t}</em>"
        if run.underline:
            t = f"<u>{t}</u>"
        text += t
    return text


def _table_to_html(table) -> str:
    rows: list[str] = ["<table>"]
    for row in table.rows:
        rows.append("<tr>")
        for cell in row.cells:
            tag = "th" if cell._element.tag == "w:th" else "td"
            cell_text = cell.text.strip().replace("\n", "<br/>")
            rows.append(f"<{tag}>{cell_text}</{tag}>")
        rows.append("</tr>")
    rows.append("</table>")
    return "\n".join(rows)


# ── XLSX ─────────────────────────────────────────────────────────────────────


async def import_xlsx(
    file: UploadFile,
    db_session,
    user_id: uuid.UUID,
    org_id: uuid.UUID,
) -> ImportResult:
    ext, content = await validate_upload(file)
    scanner = get_virus_scanner()
    scan = await scanner.scan(content, file.filename or "spreadsheet.xlsx")
    if not scan.clean:
        raise ValueError(f"Virus detected: {scan.message}")

    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheets: list[SheetInfo] = []
    html_parts: list[str] = []

    for ws_name in wb.sheetnames:
        ws = wb[ws_name]
        columns: list[str] = []
        row_count = 0
        sheet_html: list[str] = [f"<h3>{ws_name}</h3>", "<table>"]

        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                columns = [str(c) if c is not None else "" for c in row]
                sheet_html.append("<thead><tr>")
                for c in columns:
                    sheet_html.append(f"<th>{_escape_html(c)}</th>")
                sheet_html.append("</tr></thead><tbody>")
            else:
                vals = [str(v) if v is not None else "" for v in row]
                sheet_html.append("<tr>")
                for v in vals:
                    sheet_html.append(f"<td>{_escape_html(v)}</td>")
                sheet_html.append("</tr>")
            row_count += 1

        sheet_html.append("</tbody></table>")
        html_parts.extend(sheet_html)
        sheets.append(SheetInfo(name=ws_name, columns=columns, row_count=row_count))

    wb.close()
    content_html = "\n".join(html_parts)

    file_record = await _save_file(
        db_session, user_id, org_id,
        file.filename or "spreadsheet.xlsx", content, ext,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    return ImportResult(
        file_id=file_record.id,
        filename=file_record.filename,
        size_bytes=file_record.size_bytes,
        hash=file_record.hash,
        mime_type=file_record.mime_type,
        content_html=content_html,
        sheets=sheets,
    )


# ── CSV ───────────────────────────────────────────────────────────────────────


async def import_csv(
    file: UploadFile,
    db_session,
    user_id: uuid.UUID,
    org_id: uuid.UUID,
) -> ImportResult:
    ext, content = await validate_upload(file)
    scanner = get_virus_scanner()
    scan = await scanner.scan(content, file.filename or "data.csv")
    if not scan.clean:
        raise ValueError(f"Virus detected: {scan.message}")

    text = content.decode("utf-8", errors="replace")
    dialect = csv.Sniffer().sniff(text[:4096])
    reader = csv.reader(io.StringIO(text), dialect)

    html_parts: list[str] = ["<table>"]
    plain_lines: list[str] = []
    for i, row in enumerate(reader):
        tag = "th" if i == 0 else "td"
        html_parts.append("<tr>")
        for cell in row:
            escaped = _escape_html(cell)
            html_parts.append(f"<{tag}>{escaped}</{tag}>")
        html_parts.append("</tr>")
        plain_lines.append(",".join(row))
    html_parts.append("</table>")

    content_html = "\n".join(html_parts)

    file_record = await _save_file(
        db_session, user_id, org_id,
        file.filename or "data.csv", content, ext,
        "text/csv",
    )

    return ImportResult(
        file_id=file_record.id,
        filename=file_record.filename,
        size_bytes=file_record.size_bytes,
        hash=file_record.hash,
        mime_type=file_record.mime_type,
        content_html=content_html,
        plain_text="\n".join(plain_lines),
    )


# ── PDF ──────────────────────────────────────────────────────────────────────


async def import_pdf(
    file: UploadFile,
    db_session,
    user_id: uuid.UUID,
    org_id: uuid.UUID,
) -> ImportResult:
    ext, content = await validate_upload(file)
    scanner = get_virus_scanner()
    scan = await scanner.scan(content, file.filename or "document.pdf")
    if not scan.clean:
        raise ValueError(f"Virus detected: {scan.message}")

    import io as io_mod

    from pdfminer.high_level import extract_text as pdf_extract_text

    try:
        text = pdf_extract_text(io_mod.BytesIO(content))
    except Exception:
        logger.warning("PDF text extraction failed", exc_info=True)
        text = ""

    ocr_needed = len(text.strip()) < 50
    pages = max(text.count("\f") + 1, 1) if text.strip() else 1

    content_html = None
    plain_text = None
    if not ocr_needed:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        content_html = "<p>" + "</p><p>".join(
            _escape_html(p) for p in paragraphs
        ) + "</p>"
        plain_text = text.strip()

    file_record = await _save_file(
        db_session, user_id, org_id,
        file.filename or "document.pdf", content, ext,
        "application/pdf",
    )

    return ImportResult(
        file_id=file_record.id,
        filename=file_record.filename,
        size_bytes=file_record.size_bytes,
        hash=file_record.hash,
        mime_type=file_record.mime_type,
        content_html=content_html,
        plain_text=plain_text,
        pages=pages,
        ocr_needed=ocr_needed,
    )


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
