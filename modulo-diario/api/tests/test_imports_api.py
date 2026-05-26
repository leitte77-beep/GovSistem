"""Tests for the 4 import endpoints with file uploads (DOCX, XLSX, CSV, PDF)."""
import io
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.main import app
from app.core.auth import get_current_user
from app.core.database import get_db


@pytest.fixture(autouse=True)
def override_auth_and_db():
    mock_session = AsyncMock()
    mock_user = MagicMock()
    mock_user.id = uuid.uuid4()
    mock_user.organization_id = uuid.uuid4()
    mock_user.email = "autor@test.com"
    mock_user.name = "Autor"
    mock_user.is_active = True
    mock_role = MagicMock()
    mock_role.name = "ADMIN"
    mock_ur = MagicMock()
    mock_ur.role = mock_role
    mock_user.user_roles = [mock_ur]

    app.dependency_overrides[get_db] = lambda: mock_session
    app.dependency_overrides[get_current_user] = lambda: mock_user
    yield mock_session, mock_user
    app.dependency_overrides.clear()


# ── File fixtures ─────────────────────────────────────────────────────────────


@pytest.fixture
def docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("Test paragraph content.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def xlsx_bytes():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Plan1"
    ws.append(["Nome", "Valor"])
    ws.append(["Item A", "100"])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def csv_bytes():
    import csv as csv_mod
    buf = io.StringIO()
    writer = csv_mod.writer(buf)
    writer.writerow(["Produto", "Quantidade"])
    writer.writerow(["Caneta", "50"])
    return buf.getvalue().encode("utf-8")


@pytest.fixture
def pdf_bytes():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text="Invoice #12345", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


# ── DOCX Import ───────────────────────────────────────────────────────────────


@patch("app.api.v1.imports.import_docx")
@pytest.mark.anyio
async def test_upload_docx(mock_import, client, override_auth_and_db, docx_bytes):
    mock_db, mock_user = override_auth_and_db
    mock_result = MagicMock()
    mock_result.file_id = uuid.uuid4()
    mock_result.filename = "test.docx"
    mock_result.size_bytes = len(docx_bytes)
    mock_result.hash = "abc123hash"
    mock_result.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    mock_result.content_html = "<h1>Test</h1>"
    mock_result.plain_text = None
    mock_result.sheets = None
    mock_result.pages = None
    mock_result.ocr_needed = False
    mock_result.message = "File imported successfully"
    mock_import.return_value = mock_result

    files = {"file": ("test.docx", docx_bytes,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = await client.post("/api/v1/imports/docx", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "test.docx"
    assert data["file_id"] is not None


# ── XLSX Import ───────────────────────────────────────────────────────────────


@patch("app.api.v1.imports.import_xlsx")
@pytest.mark.anyio
async def test_upload_xlsx(mock_import, client, override_auth_and_db, xlsx_bytes):
    mock_db, mock_user = override_auth_and_db
    sheet = MagicMock()
    sheet.name = "Plan1"
    sheet.columns = ["Nome", "Valor"]
    sheet.row_count = 2

    mock_result = MagicMock()
    mock_result.file_id = uuid.uuid4()
    mock_result.filename = "test.xlsx"
    mock_result.size_bytes = len(xlsx_bytes)
    mock_result.hash = "hash_xlsx"
    mock_result.mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    mock_result.content_html = None
    mock_result.plain_text = None
    mock_result.sheets = [sheet]
    mock_result.pages = None
    mock_result.ocr_needed = False
    mock_result.message = "File imported successfully"
    mock_import.return_value = mock_result

    files = {"file": ("test.xlsx", xlsx_bytes,
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
    response = await client.post("/api/v1/imports/xlsx", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "test.xlsx"
    assert data["sheets"] is not None


# ── CSV Import ────────────────────────────────────────────────────────────────


@patch("app.api.v1.imports.import_csv")
@pytest.mark.anyio
async def test_upload_csv(mock_import, client, override_auth_and_db, csv_bytes):
    mock_db, mock_user = override_auth_and_db
    mock_result = MagicMock()
    mock_result.file_id = uuid.uuid4()
    mock_result.filename = "data.csv"
    mock_result.size_bytes = len(csv_bytes)
    mock_result.hash = "hash_csv"
    mock_result.mime_type = "text/csv"
    mock_result.content_html = "<table><tr><td>Produto</td></tr></table>"
    mock_result.plain_text = "Produto,Quantidade\nCaneta,50"
    mock_result.sheets = None
    mock_result.pages = None
    mock_result.ocr_needed = False
    mock_result.message = "File imported successfully"
    mock_import.return_value = mock_result

    files = {"file": ("data.csv", csv_bytes, "text/csv")}
    response = await client.post("/api/v1/imports/csv", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "data.csv"
    assert data["content_html"] is not None


# ── PDF Import ────────────────────────────────────────────────────────────────


@patch("app.api.v1.imports.import_pdf")
@pytest.mark.anyio
async def test_upload_pdf(mock_import, client, override_auth_and_db, pdf_bytes):
    mock_db, mock_user = override_auth_and_db
    mock_result = MagicMock()
    mock_result.file_id = uuid.uuid4()
    mock_result.filename = "doc.pdf"
    mock_result.size_bytes = len(pdf_bytes)
    mock_result.hash = "hash_pdf"
    mock_result.mime_type = "application/pdf"
    mock_result.content_html = "<p>Invoice #12345</p>"
    mock_result.plain_text = None
    mock_result.sheets = None
    mock_result.pages = 1
    mock_result.ocr_needed = False
    mock_result.message = "File imported successfully"
    mock_import.return_value = mock_result

    files = {"file": ("doc.pdf", pdf_bytes, "application/pdf")}
    response = await client.post("/api/v1/imports/pdf", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "doc.pdf"
    assert data["pages"] == 1
    assert data["ocr_needed"] is False
