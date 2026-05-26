"""Tests for file import endpoints (DOCX, XLSX, CSV, PDF).

Tests the conversion logic directly (no DB required).
"""

import csv
import hashlib
import io
import uuid
from unittest.mock import AsyncMock

import pytest

from app.services.pdf_utils import compute_hash
from app.services.importer import (
    import_csv,
    import_docx,
    import_pdf,
    import_xlsx,
)

TEST_USER_ID = uuid.uuid4()
TEST_ORG_ID = uuid.uuid4()


@pytest.fixture
def mock_db():
    return AsyncMock()


# ── Fixtures: create small test files ────────────────────────────────────────


@pytest.fixture
def docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Second paragraph with content.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def xlsx_bytes():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Plan1"
    ws.append(["Nome", "Valor", "Data"])
    ws.append(["Item A", "100,00", "2026-01-15"])
    ws.append(["Item B", "250,50", "2026-02-20"])
    ws2 = wb.create_sheet("Plan2")
    ws2.append(["ID", "Descrição"])
    ws2.append(["1", "Desc A"])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def csv_bytes():
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Produto", "Quantidade", "Preço"])
    writer.writerow(["Caneta", "50", "1,50"])
    writer.writerow(["Caderno", "20", "25,00"])
    writer.writerow(["Lápis", "100", "0,80"])
    return buf.getvalue().encode("utf-8")


@pytest.fixture
def pdf_bytes():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text="Invoice #12345", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 10, text="Date: 2026-05-15", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 10, text="Total: R$ 1.500,00", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


@pytest.fixture
def scanned_pdf_bytes():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, text=" ", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


# ── Helper: simulate UploadFile ──────────────────────────────────────────────


def _fake_upload(filename: str, content: bytes, mime: str):
    from unittest.mock import PropertyMock

    from fastapi import UploadFile
    file = UploadFile(filename=filename, file=io.BytesIO(content))
    type(file).content_type = PropertyMock(return_value=mime)
    return file


# ── DOCX ─────────────────────────────────────────────────────────────────────


class TestDocxImport:
    @pytest.mark.anyio
    async def test_docx_converts_to_html(self, docx_bytes, mock_db):
        file = _fake_upload("test.docx", docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        result = await import_docx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.filename == "test.docx"
        assert result.size_bytes > 0
        assert len(result.hash) == 64
        assert "<h1>" in result.content_html
        assert "<table>" in result.content_html
        assert "Test Title" in result.content_html
        assert "test paragraph" in result.content_html

    @pytest.mark.anyio
    async def test_docx_file_hash_is_sha256(self, docx_bytes, mock_db):
        file = _fake_upload("test.docx", docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        result = await import_docx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        expected = hashlib.sha256(docx_bytes).hexdigest()
        assert result.hash == expected

    @pytest.mark.anyio
    async def test_docx_table_conversion(self, docx_bytes, mock_db):
        file = _fake_upload("test.docx", docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        result = await import_docx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert "A1" in result.content_html
        assert "B2" in result.content_html


# ── XLSX ─────────────────────────────────────────────────────────────────────


class TestXlsxImport:
    @pytest.mark.anyio
    async def test_xlsx_returns_sheets(self, xlsx_bytes, mock_db):
        file = _fake_upload("test.xlsx", xlsx_bytes,
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        result = await import_xlsx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.sheets is not None
        sheet_names = [s.name for s in result.sheets]
        assert "Plan1" in sheet_names
        assert "Plan2" in sheet_names

    @pytest.mark.anyio
    async def test_xlsx_has_columns(self, xlsx_bytes, mock_db):
        file = _fake_upload("test.xlsx", xlsx_bytes,
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        result = await import_xlsx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        plan1 = [s for s in (result.sheets or []) if s.name == "Plan1"][0]
        assert "Nome" in plan1.columns
        assert "Valor" in plan1.columns
        assert plan1.row_count == 3

    @pytest.mark.anyio
    async def test_xlsx_converts_to_html_table(self, xlsx_bytes, mock_db):
        file = _fake_upload("test.xlsx", xlsx_bytes,
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        result = await import_xlsx(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert "<table>" in (result.content_html or "")
        assert "Item A" in (result.content_html or "")


# ── CSV ──────────────────────────────────────────────────────────────────────


class TestCsvImport:
    @pytest.mark.anyio
    async def test_csv_converts_to_table(self, csv_bytes, mock_db):
        file = _fake_upload("data.csv", csv_bytes, "text/csv")
        result = await import_csv(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert "<table>" in (result.content_html or "")
        assert "Produto" in (result.content_html or "")
        assert "Caneta" in (result.content_html or "")

    @pytest.mark.anyio
    async def test_csv_has_plain_text(self, csv_bytes, mock_db):
        file = _fake_upload("data.csv", csv_bytes, "text/csv")
        result = await import_csv(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.plain_text is not None
        assert "Caneta" in result.plain_text


# ── PDF ──────────────────────────────────────────────────────────────────────


class TestPdfImport:
    @pytest.mark.anyio
    async def test_pdf_with_text_extracted(self, pdf_bytes, mock_db):
        file = _fake_upload("doc.pdf", pdf_bytes, "application/pdf")
        result = await import_pdf(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.ocr_needed is False
        assert result.content_html is not None
        assert "Invoice" in result.content_html or "Invoice" in (result.plain_text or "")

    @pytest.mark.anyio
    async def test_scanned_pdf_marked_ocr(self, scanned_pdf_bytes, mock_db):
        file = _fake_upload("scan.pdf", scanned_pdf_bytes, "application/pdf")
        result = await import_pdf(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.ocr_needed is True
        assert result.content_html is None

    @pytest.mark.anyio
    async def test_pdf_has_pages(self, pdf_bytes, mock_db):
        file = _fake_upload("doc.pdf", pdf_bytes, "application/pdf")
        result = await import_pdf(file, mock_db, TEST_USER_ID, TEST_ORG_ID)
        assert result.pages >= 1


# ── Validation ───────────────────────────────────────────────────────────────


class TestFileValidation:
    def test_docx_mime_valid(self):
        from app.core.config import settings
        assert ".docx" in settings.ALLOWED_EXTENSIONS

    def test_exe_extension_not_allowed(self):
        from app.core.config import settings
        assert ".exe" not in settings.ALLOWED_EXTENSIONS

    def test_compute_hash_is_sha256(self):
        data = b"test data"
        h = compute_hash(data)
        assert len(h) == 64
        assert h == __import__("hashlib").sha256(data).hexdigest()
